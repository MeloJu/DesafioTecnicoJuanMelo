"""
Shared fixtures for E2E tests.

E2E tests require:
  - Ollama running locally with llama3.2 and nomic-embed-text loaded
  - A real photo with a visible person (YOLO can detect)

Fixture images:
  - tests/fixtures/images/worker_conforme.jpg   → person with helmet + vest
  - tests/fixtures/images/worker_nao_conforme.jpg → person without helmet

If the images are not found, the tests skip automatically.
If Ollama is unreachable, the tests skip automatically.

To run E2E tests:
    pytest tests/e2e/ -m e2e
"""
import os
import socket
import pytest
from docx import Document
from pathlib import Path

from app.rag.document_parser import DocumentParser
from app.rag.embedding_service import EmbeddingService
from app.rag.ollama_embedder import OllamaEmbedder


FIXTURES_DIR = Path(__file__).parent.parent / "fixtures"
IMAGES_DIR = FIXTURES_DIR / "images"
DOCUMENTS_DIR = FIXTURES_DIR / "documents"


# ---------------------------------------------------------------------------
# Skip guards
# ---------------------------------------------------------------------------


def _ollama_running(host: str = "127.0.0.1", port: int = 11434) -> bool:
    try:
        with socket.create_connection((host, port), timeout=2):
            return True
    except OSError:
        return False


requires_ollama = pytest.mark.skipif(
    not _ollama_running(),
    reason="Ollama não está rodando em localhost:11434 — inicie com `ollama serve`",
)


def _image_path(filename: str) -> Path:
    return IMAGES_DIR / filename


def skip_if_image_missing(filename: str):
    path = _image_path(filename)
    return pytest.mark.skipif(
        not path.exists(),
        reason=(
            f"Imagem de fixture não encontrada: {path}\n"
            "Coloque uma foto real de um trabalhador em tests/fixtures/images/"
        ),
    )


# ---------------------------------------------------------------------------
# ChromaDB + documents fixture
# ---------------------------------------------------------------------------


@pytest.fixture(scope="session")
def e2e_chroma_path(tmp_path_factory):
    """Temporary ChromaDB with construction site rules indexed."""
    db_path = tmp_path_factory.mktemp("chroma_e2e")
    return str(db_path)


@pytest.fixture(scope="session")
def indexed_fixture_documents(e2e_chroma_path):
    """
    Creates a fixture .docx with EPI rules and indexes it into the temp ChromaDB.
    Returns the chroma_path so tests can pass it to create_pipeline().
    """
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    doc_path = DOCUMENTS_DIR / "normas_obras_e2e.docx"

    # Create fixture document with construction EPI rules
    doc = Document()
    doc.add_paragraph(
        "Uso de capacete de segurança é obrigatório em todas as áreas de obra, "
        "sem exceção, conforme NR-6 e NR-18."
    )
    doc.add_paragraph(
        "Colete de alta visibilidade deve ser usado por todos os trabalhadores "
        "que circulam em áreas de movimentação de veículos."
    )
    doc.add_paragraph(
        "Botas de segurança com biqueira de aço são obrigatórias em toda a obra, "
        "conforme item 18.23.1 da NR-18."
    )
    doc.add_paragraph(
        "Luvas de proteção devem ser utilizadas em atividades com risco de corte, "
        "abrasão ou agentes químicos conforme NR-6."
    )
    doc.save(str(doc_path))

    # Index with real OllamaEmbedder (requires Ollama running)
    import chromadb

    embedder = OllamaEmbedder(model="nomic-embed-text")
    chroma_client = chromadb.PersistentClient(path=e2e_chroma_path)
    embedding_service = EmbeddingService(
        chroma_client=chroma_client,
        embedding_client=embedder,
    )

    parser = DocumentParser()
    chunks = parser.parse(str(doc_path), empresa="TestEmpresa", setor="obras")
    embedding_service.index(chunks, collection="testempresa")

    return e2e_chroma_path
